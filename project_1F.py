# -*- coding: utf-8 -*-
"""
Created on Fri Dec  4 16:00:51 2020

@author: Vikas
"""

import PowerController
import CANOE_Class
import GUI
from PyQt5 import QtCore, QtGui, QtWidgets
import time
from docx import Document
from docx.shared import Inches
import os

power = PowerController.power_control
CANapp = CANOE_Class.CANoe()
document = Document()
document.add_heading('Test Report', 3)
p1 = document.add_paragraph("Software version being used for the test is ")
p1.add_run(str(CANapp.sof_ver)).bold = True
fname = 'AutoGenartedTestReport.docx'
try :
    document.save(fname)
except:
    res_check = 0
    while res_check == 0: 
        req = input("Old report is open, do yo want to override (Y/N?): ")
        if req == "Y" or req == "y":
            print("Please close the current report to acess for overwrite.")
            time.sleep(2)
            check_doc_close = 0
            while check_doc_close == 0:
                try:
                    document.save(fname)
                    check_doc_close = 1
                except :
                    print("Document not closed properly.")
                    time.sleep(2)
                    check_doc_close = 0
            res_check = 1        
        elif req == "N" or req == "n":
            fname = input("Enter the new name for the report: ")
            document.save(fname)
            res_check = 1
        else :
            print("Response received is not n Y or N format")
            res_check = 0

if __name__ == "__main__":
    ## OPenign GUI to control the Power supply
    import sys
       
    ## Acessing CANoe application once the power supply setup is complete
    file_name = 'pr1_test_automation'                           # Provide the configuration file name to open
    CANapp.open_cfg(file_name)  
    p2 = document.add_paragraph("Configuration file being used is ")
    p2.add_run(file_name).italic = True
    ## Getting the list of al system variable and theie value for a specific namespace
    name_space = 'TestCase'
    p3 = document.add_paragraph('Below table provide details for System variable in name space ')
    p3.add_run(name_space).bold = True
    p3.add_run(' with initial values and changed values through python')
    table1 = document.add_table(rows=1, cols=4)
    hdr_cells_t1 = table1.rows[0].cells
    hdr_cells_t1[0].text = 'Variable Name'
    hdr_cells_t1[1].text = 'Initial Value'
    hdr_cells_t1[2].text = 'Before test Reset to Value'
    hdr_cells_t1[3].text = 'After test Value'
    sys_var_detail = CANapp.get_allsysVar(name_space)
    for var_name, var_value in sys_var_detail:
        row_cells_t1 = table1.add_row().cells
        row_cells_t1[0].text = var_name
        row_cells_t1[1].text = str(var_value)
    # print(sys_var_detail)
        
    ## Information for the message and signal value from the current dbc file
    eng_msg_sig_data = [('EngineData','EngineRPM'),('EngineData','EngineTemp'),('EngineData','EngineStatus')]
    power_msg_sig_data = [('PowerMode','CarMode'),('PowerMode','PowerMode')]
    
    ## Starting the simulation in the CANoe
    CANapp.start_measure()  
    # time.sleep(10)
    
    ## Setting the system variable value in real time
    for i in sys_var_detail:
        CANapp.set_sysVar(name_space, i[0], float(0))
    
    ## To get the value if the the system variable after change
    sys_var_change = CANapp.get_allsysVar(name_space)
    hdr_cells_t1 = table1.columns[2].cells
    for i in range(len(sys_var_change)):
        hdr_cells_t1[i+1].text = str(sys_var_change[i][1])
    # for i in sys_var_detail:
    #     print("Value of the system variable",i[0]," after change is",CANapp.get_sysVar(name_space, i[0]))
    
    ## printing out the signal 
    document.add_paragraph('Below table provide the signal details.')
    # p4.add_run(eng_msg_sig_data[0][0]).bold = True
    table2 = document.add_table(rows=1, cols=5)
    hdr_cells_t2 = table2.rows[0].cells
    hdr_cells_t2[0].text = 'CAN Channel'
    hdr_cells_t2[1].text = 'Message Name'
    hdr_cells_t2[2].text = 'Signal Name'
    hdr_cells_t2[3].text = 'Signal Value'
    hdr_cells_t2[4].text = 'After test Value'
    for i in eng_msg_sig_data:
        sig_val = CANapp.get_signal(1, i[0], i[1])
        row_cells_t2 = table2.add_row().cells
        row_cells_t2[0].text = '1'
        row_cells_t2[1].text = i[0]
        row_cells_t2[2].text = i[1]
        row_cells_t2[3].text = str(sig_val)
        # print("Signal from message",i[0],"with name",i[1],"have value",sig_val)
    
    for i in power_msg_sig_data:
        sig_val = CANapp.get_signal(1, i[0], i[1])
        row_cells_t2 = table2.add_row().cells
        row_cells_t2[0].text = '1'
        row_cells_t2[1].text = i[0]
        row_cells_t2[2].text = i[1]
        row_cells_t2[3].text = str(sig_val)
        # print("Signal from message",i[0],"with name",i[1],"have value",sig_val)

    document.add_page_break()
    
    ## Getting the list of all test enviornment and test module available
    test_env_list = CANapp.get_list_tes_env()
    p4 = document.add_paragraph('There are total ')
    p4.add_run(str(len(test_env_list))).bold = True
    p4.add_run(' test enviornments with name/s ')
    for i in test_env_list:
        p4.add_run(i+' ').italic = True
    test_detail = []
    for i in test_env_list:
        val = CANapp.list_test_module(i)
        # p4.add_run(val).italic = True
        test_detail.append((i,val))
    # print(test_detail)
    
    ## For loop details to acess the required variabkes and perform the required function 
    # for i in test_detail:
    #     for j in i[1]:
    #         print(i[0] + ' ' + j)
            
    ## To acess specific value in test details use the format as below
    # for first tuple test_detail[0]
    # for the test enviornmrnt name from the first value tuple test_detail[0][0]
    # to acess the first test modulefrom the first test enviornmrnt test_detail[0][1][0]
    ## Getting the list and values of all system varoables available in the system
    
    document.add_paragraph('Below table provide the result for the testmodules results.')
    table3 = document.add_table(rows=1, cols=4)
    hdr_cells_t3 = table3.rows[0].cells
    hdr_cells_t3[0].text = 'Test Env/ Name'
    hdr_cells_t3[1].text = 'Test Module'
    hdr_cells_t3[2].text = 'Test Case Number'
    hdr_cells_t3[3].text = 'Result'
    ## Run the test scripts from the enviornment
    for i in test_detail:
        for j in i[1]:
            result = CANapp.test_module_run(i[0], j)
            # print(result)
            count = 0
            for check in result[0]:
                if check == 1:
                    row_cells_t3 = table3.add_row().cells
                    count = count + 1
                    row_cells_t3[0].text = i[0]
                    row_cells_t3[1].text = j
                    row_cells_t3[2].text = str(count)
                    row_cells_t3[3].text = 'Pass'
                    # print("From the test enviornment",i[0],"for test module",j,"test case number",count,"is pass")
                else:
                    row_cells_t2 = table2.add_row().cells
                    count = count + 1
                    row_cells_t3[0].text = i[0]
                    row_cells_t3[1].text = j
                    row_cells_t3[2].text = str(count)
                    row_cells_t3[3].text = 'Fail'
                    # print("From the test enviornment",i[0],"for test module",j,"test case number",count,"is fail")
    
    ## To check value of system variable after test case run
    for i in sys_var_detail:
        print("Value of the system variable",i[0]," with chnges done during test case with CAPL script is",CANapp.get_sysVar(name_space, i[0]))
    ## To get the value if the the system variable after change
    sys_var_aftest = CANapp.get_allsysVar(name_space)
    hdr_cells_t1 = table1.columns[3].cells
    for i in range(len(sys_var_aftest)):
        hdr_cells_t1[i+1].text = str(sys_var_aftest[i][1])
    
    hdr_cells_t2 = table2.columns[4].cells
    ## printing out the signal information
    for i in range(len(eng_msg_sig_data)):
        sig_val = CANapp.get_signal(1, eng_msg_sig_data[i][0], eng_msg_sig_data[i][1])
        hdr_cells_t2[i+1].text = str(sig_val)
        # print("Signal from message",i[0],"with name",i[1],"have value",sig_val,"after changed with CAPL script in test module")
    
    for i in range(len(power_msg_sig_data)):
        sig_val = CANapp.get_signal(1, power_msg_sig_data[i][0], power_msg_sig_data[i][1])
        hdr_cells_t2[i+1+len(eng_msg_sig_data)].text = str(sig_val)
        # print("Signal from message",i[0],"with name",i[1],"have value",sig_val,"after changed with CAPL script in test module")        
    
    ## Stop the measurement
    CANapp.stop_measure()
    document.save(fname)
    file_path = os.path.abspath(fname)
    os.startfile(file_path)
    CANapp.Close()
    app = QtWidgets.QApplication(sys.argv)
    Dialog = QtWidgets.QDialog()
    ui = GUI.Ui_Dialog()
    ui.setupUi(Dialog)
    Dialog.show()
    sys.exit(app.exec_())